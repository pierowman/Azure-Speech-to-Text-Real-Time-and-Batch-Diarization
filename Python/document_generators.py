"""
Document generation utilities for creating Word documents and other exports.
"""
from datetime import datetime
from typing import List, Dict, Any
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import logging

logger = logging.getLogger(__name__)


class AuditLogDocumentGenerator:
    """Generator for audit log Word documents."""
    
    @staticmethod
    def create_document(audit_log: List[Dict[str, Any]]) -> Document:
        """
        Create a Word document from audit log entries.
        
        Args:
            audit_log: List of audit log entry dictionaries
            
        Returns:
            Document object ready to be saved
        """
        doc = Document()
        
        # Add title
        title = doc.add_heading('Transcription Edit Audit Log', level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add separator
        doc.add_paragraph('=' * 80)
        
        # Add metadata
        doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        doc.add_paragraph(f"Total Edits: {len(audit_log)}")
        doc.add_paragraph()
        
        # Add edits section
        doc.add_heading('Transcription Edits', level=2)
        doc.add_paragraph('Edits organized by segment position in the transcript:')
        doc.add_paragraph('-' * 80)
        
        # Sort edits by segment order
        edits_by_segment = sorted(
            audit_log,
            key=lambda x: (x.get('lineNumber', 999999), x.get('timestamp', ''))
        )
        
        # Add each edit
        for i, edit in enumerate(edits_by_segment, 1):
            AuditLogDocumentGenerator._add_edit_entry(doc, i, edit)
            doc.add_paragraph()  # Add spacing
        
        return doc
    
    @staticmethod
    def _add_edit_entry(doc: Document, index: int, edit: Dict[str, Any]) -> None:
        """
        Add a single edit entry to the document.
        
        Args:
            doc: Document to add to
            index: Edit number
            edit: Edit entry dictionary
        """
        action = edit.get('action', 'edit')
        
        if action == 'bulk_speaker_rename':
            AuditLogDocumentGenerator._add_bulk_speaker_rename(doc, index, edit)
        elif action == 'bulk_speaker_reassignment':
            AuditLogDocumentGenerator._add_bulk_speaker_reassignment(doc, index, edit)
        else:
            AuditLogDocumentGenerator._add_regular_edit(doc, index, edit)
    
    @staticmethod
    def _add_bulk_speaker_rename(doc: Document, index: int, edit: Dict[str, Any]) -> None:
        """Add bulk speaker rename entry."""
        p = doc.add_paragraph()
        p.add_run(f"Operation #{index} - Bulk Speaker Rename").bold = True
        p.runs[-1].font.size = Pt(11)
        
        info_p = doc.add_paragraph(style='List Bullet')
        info_p.add_run(f"Time: {edit.get('timestamp', 'N/A')}\n")
        info_p.add_run(f"Action: Bulk Speaker Rename\n")
        info_p.add_run(f"Description: {edit.get('description', 'N/A')}")
        
        rename_p = doc.add_paragraph(style='List Bullet')
        rename_p.add_run(f"Speaker Renamed: ").bold = True
        rename_p.add_run(f'"{edit.get("oldSpeaker", "")}" -> "{edit.get("newSpeaker", "")}"')
        rename_p.runs[-1].font.color.rgb = RGBColor(102, 126, 234)
        
        count_p = doc.add_paragraph(style='List Bullet')
        count_p.add_run(f"Affected Segments: {edit.get('segmentCount', 0)} segment(s)")
    
    @staticmethod
    def _add_bulk_speaker_reassignment(doc: Document, index: int, edit: Dict[str, Any]) -> None:
        """Add bulk speaker reassignment entry."""
        p = doc.add_paragraph()
        p.add_run(f"Operation #{index} - Bulk Speaker Reassignment").bold = True
        p.runs[-1].font.size = Pt(11)
        
        info_p = doc.add_paragraph(style='List Bullet')
        info_p.add_run(f"Time: {edit.get('timestamp', 'N/A')}\n")
        info_p.add_run(f"Action: Bulk Speaker Reassignment\n")
        info_p.add_run(f"Description: {edit.get('description', 'N/A')}")
        
        reassign_p = doc.add_paragraph(style='List Bullet')
        reassign_p.add_run(f"Reassignment: ").bold = True
        reassign_p.add_run(f'"{edit.get("fromSpeaker", "")}" -> "{edit.get("toSpeaker", "")}"')
        reassign_p.runs[-1].font.color.rgb = RGBColor(102, 126, 234)
        
        count_p = doc.add_paragraph(style='List Bullet')
        count_p.add_run(f"Affected Segments: {edit.get('segmentCount', 0)} segment(s)")
    
    @staticmethod
    def _add_regular_edit(doc: Document, index: int, edit: Dict[str, Any]) -> None:
        """Add regular edit entry."""
        action = edit.get('action', 'edit')
        
        p = doc.add_paragraph()
        p.add_run(f"Edit #{index} - Segment #{edit.get('lineNumber', 'N/A')}").bold = True
        p.runs[-1].font.size = Pt(11)
        
        # Timestamp and action
        info_p = doc.add_paragraph(style='List Bullet')
        info_p.add_run(f"Time: {edit.get('timestamp', 'N/A')}\n")
        info_p.add_run(f"Action: {action}\n")
        info_p.add_run(f"Speaker: {edit.get('speaker', 'Unknown')}\n")
        info_p.add_run(f"Audio Time: {edit.get('startTime', '0:00')}")
        
        # Speaker change details if applicable
        if 'oldSpeaker' in edit and 'newSpeaker' in edit:
            speaker_p = doc.add_paragraph(style='List Bullet')
            speaker_p.add_run(f"Speaker Change: ").bold = True
            speaker_p.add_run(f"{edit.get('oldSpeaker', '')} -> {edit.get('newSpeaker', '')}")
            speaker_p.runs[-1].font.color.rgb = RGBColor(102, 126, 234)
        
        # Text changes
        if edit.get('oldText') != edit.get('newText'):
            old_p = doc.add_paragraph(style='List Bullet')
            old_p.add_run("Old Text: ").bold = True
            old_p.add_run(edit.get('oldText', ''))
            old_p.runs[-1].font.color.rgb = RGBColor(220, 53, 69)
            
            new_p = doc.add_paragraph(style='List Bullet')
            new_p.add_run("New Text: ").bold = True
            new_p.add_run(edit.get('newText', ''))
            new_p.runs[-1].font.color.rgb = RGBColor(0, 200, 81)


class TranscriptionDocumentGenerator:
    """Generator for transcription Word documents."""
    
    @staticmethod
    def create_document(segments: List[Dict[str, Any]]) -> Document:
        """
        Create a formatted Word document from transcription segments.
        
        Args:
            segments: List of segment dictionaries
            
        Returns:
            Document object ready to be saved
        """
        doc = Document()
        
        # Add title
        title = doc.add_heading('Speech to Text Transcription with Diarization', level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add separator
        doc.add_paragraph('=' * 50)
        
        # Add metadata
        doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        doc.add_paragraph()
        
        # Add transcript header
        doc.add_heading('TRANSCRIPT:', level=2)
        doc.add_paragraph('-' * 50)
        doc.add_paragraph()
        
        # Add segments
        for seg_data in segments:
            TranscriptionDocumentGenerator._add_segment(doc, seg_data)
        
        return doc
    
    @staticmethod
    def _add_segment(doc: Document, seg_data: Dict[str, Any]) -> None:
        """
        Add a single segment to the document.
        
        Args:
            doc: Document to add to
            seg_data: Segment data dictionary
        """
        line_num = seg_data.get('lineNumber', 0)
        speaker = seg_data.get('speaker', 'Unknown')
        text = seg_data.get('text', '')
        timestamp = seg_data.get('uiFormattedStartTime', '')
        
        # Speaker line with timestamp
        p = doc.add_paragraph()
        p.add_run(f"#{line_num} ").bold = True
        p.runs[-1].font.color.rgb = RGBColor(128, 128, 128)
        p.add_run(f"[{timestamp}] {speaker}:").bold = True
        
        # Text paragraph (indented)
        text_p = doc.add_paragraph(text)
        text_p.paragraph_format.left_indent = Pt(36)
        text_p.paragraph_format.space_after = Pt(6)


class CombinedDocumentGenerator:
    """Generator for combined transcription and audit log Word documents."""
    
    @staticmethod
    def create_document(segments: List[Dict[str, Any]], audit_log: List[Dict[str, Any]]) -> Document:
        """
        Create a combined Word document with both transcription and audit log.
        
        Args:
            segments: List of segment dictionaries
            audit_log: List of audit log entry dictionaries
            
        Returns:
            Document object ready to be saved
        """
        doc = Document()
        
        # Add main title
        title = doc.add_heading('Speech to Text Transcription with Edit History', level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add separator
        doc.add_paragraph('=' * 80)
        
        # Add metadata
        doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        doc.add_paragraph(f"Total Segments: {len(segments)}")
        doc.add_paragraph(f"Total Edits: {len(audit_log)}")
        doc.add_paragraph()
        
        # ? Build set of individually edited segment indices (same logic as frontend)
        individually_edited_segments = set()
        individual_actions = ['segment_edit', 'speaker_change', 'edit_with_speaker_change', 'edit']
        
        for entry in audit_log:
            if entry.get('action') in individual_actions and entry.get('segmentIndex') is not None:
                individually_edited_segments.add(entry.get('segmentIndex'))
        
        logger.info(f"?? Document generation: Found {len(individually_edited_segments)} individually edited segments")
        
        # Add transcript section
        doc.add_heading('PART 1: TRANSCRIBED TEXT', level=2)
        doc.add_paragraph('-' * 80)
        doc.add_paragraph()
        
        for idx, seg_data in enumerate(segments):
            # Check if this segment was individually edited (not bulk operation)
            is_edited = idx in individually_edited_segments
            CombinedDocumentGenerator._add_segment(doc, seg_data, is_edited)
        
        # Add page break before audit log
        doc.add_page_break()
        
        # Add audit log section
        doc.add_heading('PART 2: EDIT AUDIT LOG', level=2)
        doc.add_paragraph('-' * 80)
        doc.add_paragraph()
        
        if audit_log and len(audit_log) > 0:
            doc.add_paragraph('Edits organized by segment position in the transcript:')
            doc.add_paragraph('-' * 80)
            
            # Sort edits by segment order
            edits_by_segment = sorted(
                audit_log,
                key=lambda x: (x.get('lineNumber', 999999), x.get('timestamp', ''))
            )
            
            # Add each edit
            for i, edit in enumerate(edits_by_segment, 1):
                CombinedDocumentGenerator._add_edit_entry(doc, i, edit)
                doc.add_paragraph()  # Add spacing
        else:
            doc.add_paragraph('No edits have been made to this transcription.')
        
        return doc
    
    @staticmethod
    def _add_segment(doc: Document, seg_data: Dict[str, Any], is_edited: bool = False) -> None:
        """
        Add a single segment to the document.
        
        Args:
            doc: Document to add to
            seg_data: Segment data dictionary
            is_edited: Whether this segment was individually edited (determined from audit log)
        """
        line_num = seg_data.get('lineNumber', 0)
        speaker = seg_data.get('speaker', 'Unknown')
        text = seg_data.get('text', '')
        timestamp = seg_data.get('uiFormattedStartTime', '')
        
        # Speaker line with timestamp
        p = doc.add_paragraph()
        p.add_run(f"#{line_num} ").bold = True
        p.runs[-1].font.color.rgb = RGBColor(128, 128, 128)
        p.add_run(f"[{timestamp}] {speaker}:").bold = True
        
        # ? Add edited indicator ONLY if segment was individually edited (not bulk operation)
        if is_edited:
            p.add_run(" [EDITED]").bold = True
            p.runs[-1].font.color.rgb = RGBColor(0, 200, 81)
        
        # Text paragraph (indented)
        text_p = doc.add_paragraph(text)
        text_p.paragraph_format.left_indent = Pt(36)
        text_p.paragraph_format.space_after = Pt(6)
    
    @staticmethod
    def _add_edit_entry(doc: Document, index: int, edit: Dict[str, Any]) -> None:
        """Add a single edit entry to the document."""
        action = edit.get('action', 'edit')
        
        if action == 'bulk_speaker_rename':
            CombinedDocumentGenerator._add_bulk_speaker_rename(doc, index, edit)
        elif action == 'bulk_speaker_reassignment':
            CombinedDocumentGenerator._add_bulk_speaker_reassignment(doc, index, edit)
        elif action == 'bulk_speaker_delete':
            CombinedDocumentGenerator._add_bulk_speaker_delete(doc, index, edit)
        else:
            CombinedDocumentGenerator._add_regular_edit(doc, index, edit)
    
    @staticmethod
    def _add_bulk_speaker_rename(doc: Document, index: int, edit: Dict[str, Any]) -> None:
        """Add bulk speaker rename entry."""
        p = doc.add_paragraph()
        p.add_run(f"Operation #{index} - Bulk Speaker Rename").bold = True
        p.runs[-1].font.size = Pt(11)
        
        info_p = doc.add_paragraph(style='List Bullet')
        info_p.add_run(f"Time: {edit.get('timestamp', 'N/A')}\n")
        info_p.add_run(f"Action: Bulk Speaker Rename\n")
        info_p.add_run(f"Description: {edit.get('description', 'N/A')}")
        
        rename_p = doc.add_paragraph(style='List Bullet')
        rename_p.add_run(f"Speaker Renamed: ").bold = True
        rename_p.add_run(f'"{edit.get("oldSpeaker", "")}" -> "{edit.get("newSpeaker", "")}"')
        rename_p.runs[-1].font.color.rgb = RGBColor(102, 126, 234)
        
        count_p = doc.add_paragraph(style='List Bullet')
        count_p.add_run(f"Affected Segments: {edit.get('segmentCount', 0)} segment(s)")
    
    @staticmethod
    def _add_bulk_speaker_reassignment(doc: Document, index: int, edit: Dict[str, Any]) -> None:
        """Add bulk speaker reassignment entry."""
        p = doc.add_paragraph()
        p.add_run(f"Operation #{index} - Bulk Speaker Reassignment").bold = True
        p.runs[-1].font.size = Pt(11)
        
        info_p = doc.add_paragraph(style='List Bullet')
        info_p.add_run(f"Time: {edit.get('timestamp', 'N/A')}\n")
        info_p.add_run(f"Action: Bulk Speaker Reassignment\n")
        info_p.add_run(f"Description: {edit.get('description', 'N/A')}")
        
        reassign_p = doc.add_paragraph(style='List Bullet')
        reassign_p.add_run(f"Reassignment: ").bold = True
        reassign_p.add_run(f'"{edit.get("fromSpeaker", "")}" -> "{edit.get("toSpeaker", "")}"')
        reassign_p.runs[-1].font.color.rgb = RGBColor(102, 126, 234)
        
        count_p = doc.add_paragraph(style='List Bullet')
        count_p.add_run(f"Affected Segments: {edit.get('segmentCount', 0)} segment(s)")
    
    @staticmethod
    def _add_bulk_speaker_delete(doc: Document, index: int, edit: Dict[str, Any]) -> None:
        """Add bulk speaker delete entry."""
        p = doc.add_paragraph()
        p.add_run(f"Operation #{index} - Bulk Speaker Delete").bold = True
        p.runs[-1].font.size = Pt(11)
        
        info_p = doc.add_paragraph(style='List Bullet')
        info_p.add_run(f"Time: {edit.get('timestamp', 'N/A')}\n")
        info_p.add_run(f"Action: Bulk Speaker Delete\n")
        info_p.add_run(f"Description: {edit.get('description', 'N/A')}")
        
        delete_p = doc.add_paragraph(style='List Bullet')
        delete_p.add_run(f"Speaker Deleted: ").bold = True
        delete_p.add_run(f'"{edit.get("oldSpeaker", "")}" -> "{edit.get("newSpeaker", "")}"')
        delete_p.runs[-1].font.color.rgb = RGBColor(102, 126, 234)
        
        count_p = doc.add_paragraph(style='List Bullet')
        count_p.add_run(f"Affected Segments: {edit.get('segmentCount', 0)} segment(s)")
    
    @staticmethod
    def _add_regular_edit(doc: Document, index: int, edit: Dict[str, Any]) -> None:
        """Add regular edit entry."""
        action = edit.get('action', 'edit')
        
        p = doc.add_paragraph()
        p.add_run(f"Edit #{index} - Segment #{edit.get('lineNumber', 'N/A')}").bold = True
        p.runs[-1].font.size = Pt(11)
        
        # Timestamp and action
        info_p = doc.add_paragraph(style='List Bullet')
        info_p.add_run(f"Time: {edit.get('timestamp', 'N/A')}\n")
        info_p.add_run(f"Action: {action}\n")
        info_p.add_run(f"Speaker: {edit.get('speaker', 'Unknown')}\n")
        info_p.add_run(f"Audio Time: {edit.get('startTime', '0:00')}")
        
        # Speaker change details if applicable
        if 'oldSpeaker' in edit and 'newSpeaker' in edit and edit.get('oldSpeaker') != edit.get('newSpeaker'):
            speaker_p = doc.add_paragraph(style='List Bullet')
            speaker_p.add_run(f"Speaker Change: ").bold = True
            speaker_p.add_run(f"{edit.get('oldSpeaker', '')} -> {edit.get('newSpeaker', '')}")
            speaker_p.runs[-1].font.color.rgb = RGBColor(102, 126, 234)
        
        # Text changes (only show if actually different)
        if edit.get('oldText') != edit.get('newText'):
            old_p = doc.add_paragraph(style='List Bullet')
            old_p.add_run("Old Text: ").bold = True
            old_p.add_run(edit.get('oldText', ''))
            old_p.runs[-1].font.color.rgb = RGBColor(220, 53, 69)
            
            new_p = doc.add_paragraph(style='List Bullet')
            new_p.add_run("New Text: ").bold = True
            new_p.add_run(edit.get('newText', ''))
            new_p.runs[-1].font.color.rgb = RGBColor(0, 200, 81)
